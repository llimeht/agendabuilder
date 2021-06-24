"""
Representation of the agenda summary given to attendees of a meeting
"""

from pathlib import Path
from typing import (
    Optional,
    Union,
)

import docx  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
import docx.table  # type: ignore

from .meeting import Agenda, AgendaItem, AgendaHeading
from .locator import FileLocator


class AgendaListing:
    def __init__(
        self,
        meeting: Agenda,
        template: Union[str, Path],
        locator: Optional[FileLocator] = None,
    ) -> None:
        self.meeting = meeting
        self.template = template
        self.locator = locator or Path

        self.star = "🟊 "
        self.star_font = "Symbola"
        self.agenda_heading_style = "AgendaHeading"
        self.agenda_item_style = "AgendaItem"
        self.agenda_heading_bg_color = "DDDDDD"
        self.document: docx.Document = None

    @staticmethod
    def _set_cell_bg_color(
        cell: docx.table._Cell,  # pylint: disable=protected-access
        color: str,
    ) -> docx.table._Cell:  # pylint: disable=protected-access
        """
        set background shading for Header Rows
        """
        tc = cell._tc  # pylint: disable=protected-access
        props = tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        props.append(shading)
        return cell

    def load_template(self) -> None:
        resolved_filename = self.locator(self.template)
        self.document = docx.Document(resolved_filename)

    def find_table(self) -> docx.table.Table:
        return self.document.tables[0]

    def _fill_header(
        self,
        item: AgendaHeading,
        row: docx.table._Row,  # pylint: disable=protected-access
    ) -> None:
        cells = row.cells

        # merge the cells together
        merged_cell = cells[1].merge(cells[2])
        merged_cell = merged_cell.merge(cells[3])

        # add title
        merged_cell.text = item.title

        # add styling
        merged_cell.paragraphs[0].style = self.agenda_heading_style
        cells[0].paragraphs[0].style = self.agenda_heading_style
        self._set_cell_bg_color(merged_cell, self.agenda_heading_bg_color)
        self._set_cell_bg_color(cells[0], self.agenda_heading_bg_color)

    def _fill_item(
        self, item: AgendaItem, row: docx.table._Row  # pylint: disable=protected-access
    ) -> None:
        cells = row.cells

        # add details
        if item.starred:
            cells[1].text = ""
            p = cells[1].paragraphs[0]
            r = p.add_run(self.star)
            r.font.name = "Symbola"
            p.add_run(item.title)
        else:
            cells[1].text = item.title
        if item.action:
            cells[2].text = item.action
        if item.who:
            cells[3].text = item.who

        # add styling
        for c in cells:
            c.paragraphs[0].style = self.agenda_item_style

    def fill_table(self, table: docx.table.Table) -> None:
        for item in self.meeting.items:
            itemnum = str(item.num)
            print("  Adding %s" % itemnum)

            if isinstance(item, AgendaHeading):
                row = table.add_row()
                row.cells[0].text = itemnum
                self._fill_header(item, row)
            elif isinstance(item, AgendaItem):
                row = table.add_row()
                row.cells[0].text = itemnum
                self._fill_item(item, row)

    def build(self) -> None:
        self.load_template()
        table = self.find_table()
        self.fill_table(table)

    def save(self, filename: Union[str, Path]) -> None:
        resolved_filename = self.locator(filename)
        self.document.save(resolved_filename)
